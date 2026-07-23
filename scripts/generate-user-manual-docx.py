from __future__ import annotations

import importlib.util
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PDF_GENERATOR = ROOT / "scripts" / "generate-user-manual-full.py"
OUTPUT_DIR = ROOT / "output" / "docx"
OUTPUT = OUTPUT_DIR / "STCR-User-Manual-Source.docx"
ARCHIVE = OUTPUT_DIR / "archive" / "STCR-User-Manual-Source-2026-07-23.docx"
ANNOTATED_DIR = ROOT / "output" / "manual-assets" / "annotated"

# compact_reference_guide, with named STCR overrides for A4 and Thai typography.
PAGE_WIDTH_DXA = 11906
PAGE_HEIGHT_DXA = 16838
CONTENT_WIDTH_DXA = 10205
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "202733"
INK = "27313D"
MUTED = "647181"
YELLOW = "F1CE24"
PALE_YELLOW = "FFFBE8"
RED = "E1262F"
LINE = "D8DEE6"
WHITE = "FFFFFF"
LIGHT = "F3F5F7"


def load_manual_module():
    spec = importlib.util.spec_from_file_location("stcr_manual_pdf", PDF_GENERATOR)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load {PDF_GENERATOR}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin not in kwargs:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(kwargs[margin]))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS_DXA)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=9, bold=False, color=INK, italic=False):
    run.font.name = "Sarabun"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Sarabun")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Sarabun")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Sarabun")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_tokens(paragraph, *, before=0, after=4, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_rich_text(paragraph, text: str, *, size=9, color=INK):
    parts = re.split(r"(<b>.*?</b>)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("<b>") and part.endswith("</b>")
        clean = re.sub(r"</?b>", "", part)
        set_run_font(paragraph.add_run(clean), size=size, bold=bold, color=color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=8, color=MUTED)


def configure_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.right_margin = Mm(15)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(15)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_tokens(p, before=0, after=0, line=1)
    set_run_font(
        p.add_run("STCR  |  Smoking Temperature Control"),
        size=8.5,
        bold=True,
        color=NAVY,
    )
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), YELLOW)
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Mm(180))
    set_table_geometry(table, [7800, 2405], indent_dxa=0)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    set_paragraph_tokens(left, after=0, line=1)
    set_run_font(left.add_run("คู่มือการใช้งานระบบ STCR"), size=8, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_tokens(right, after=0, line=1)
    set_run_font(right.add_run("หน้า "), size=8, color=MUTED)
    add_page_field(right)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Sarabun"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Sarabun")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Sarabun")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Sarabun")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Sarabun"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Sarabun")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Sarabun")
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Sarabun")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(NAVY)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Sarabun"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Sarabun")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Sarabun")
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Sarabun")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True


def add_title(document, text, *, size, color=NAVY, after=4):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, before=0, after=after, line=1.0)
    set_run_font(p.add_run(text), size=size, bold=True, color=color)
    return p


def add_body(document, text, *, size=9, color=INK, align=None, after=4, bold=False):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_tokens(p, after=after, line=1.25)
    set_run_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_callout(document, text):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_YELLOW)
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    set_paragraph_tokens(p, after=0, line=1.2)
    add_rich_text(p, text, size=8.4, color=NAVY)


def add_annotated_image(document, filename):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_tokens(p, after=4, line=1)
    p.add_run().add_picture(str(ANNOTATED_DIR / filename), width=Mm(178))


def fill_step_cell(cell, number, text):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=80, bottom=80, start=80, end=100)
    p = cell.paragraphs[0]
    set_paragraph_tokens(p, after=0, line=1.2)
    badge = p.add_run(f" {number} ")
    set_run_font(badge, size=8.5, bold=True, color=WHITE)
    badge_props = badge._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), RED)
    badge_props.append(shading)
    set_run_font(p.add_run("  "), size=8.5, color=INK)
    add_rich_text(p, text, size=8.5)


def add_steps(document, steps):
    rows = (len(steps) + 1) // 2
    table = document.add_table(rows=rows, cols=2)
    set_table_geometry(table, [5103, 5102])
    index = 0
    for row in table.rows:
        for cell in row.cells:
            if index < len(steps):
                number, text = steps[index]
                fill_step_cell(cell, number, text)
            else:
                cell.text = ""
            index += 1
    return table


def add_section_page(document, section):
    p = document.add_paragraph(style="Heading 1")
    set_run_font(p.add_run(section["title"]), size=16, bold=True, color=NAVY)

    intro = document.add_paragraph()
    set_paragraph_tokens(intro, after=5, line=1.2)
    set_run_font(intro.add_run(section["intro"]), size=8.7, color=MUTED)

    add_annotated_image(document, section["image"])
    add_steps(document, section["steps"])
    add_callout(document, section["note"])
    document.add_page_break()


def add_troubleshooting(document):
    title = document.add_paragraph(style="Heading 1")
    set_run_font(title.add_run("11. การแก้ปัญหาเบื้องต้น"), size=16, bold=True, color=NAVY)
    add_body(
        document,
        "ตรวจตามรายการนี้ก่อนแจ้งผู้ดูแลระบบ พร้อมระบุบริษัท หมายเลขเตา รอบ และเวลาที่พบปัญหา",
        size=8.8,
        color=MUTED,
        after=6,
    )
    rows = [
        ("อาการ", "สิ่งที่ต้องตรวจ"),
        ("Failed to fetch", "ตรวจ Node-RED API, Tunnel, MySQL และเครือข่าย แล้วกดลองใหม่"),
        ("ค่าไม่เปลี่ยน", "ตรวจเวลาอัปเดตล่าสุด สถานะ MQTT และ Topic ของบริษัท/เตา"),
        ("ขาดการเชื่อมต่อ", "ตรวจว่าต้นทางยังส่งข้อมูลและเครื่องโรงงานออนไลน์"),
        ("รายงานไม่ตรงกราฟ", "ตรวจบริษัท เตา รอบ ช่วงเวลา แล้วกดโหลดพรีวิวใหม่"),
        ("ดาวน์โหลดไม่ได้", "อนุญาตการดาวน์โหลดของเบราว์เซอร์ และตรวจพื้นที่จัดเก็บ"),
        ("หน้าเว็บยังเป็นแบบเดิม", "กด Ctrl + F5 เพื่อโหลดไฟล์เว็บเวอร์ชันล่าสุด"),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 7505])
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_shading(cell, NAVY if row_index == 0 else (LIGHT if row_index % 2 == 0 else WHITE))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            set_paragraph_tokens(p, after=0, line=1.2)
            set_run_font(
                p.add_run(text),
                size=8.8,
                bold=row_index == 0 or (row_index > 0 and col_index == 0),
                color=WHITE if row_index == 0 else INK,
            )
    add_callout(
        document,
        "<b>ข้อมูลที่ควรส่งให้ผู้ดูแลเมื่อแจ้งปัญหา</b><br/>"
        "บริษัท GR หรือ TTN, หมายเลขเตา, หมายเลขรอบ, วันเวลา และภาพหน้าจอที่พบปัญหา",
    )


def build_docx():
    manual = load_manual_module()
    manual.annotate_screenshots()

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    document.core_properties.title = "คู่มือการใช้งานระบบ STCR"
    document.core_properties.subject = "Smoking Temperature Control"
    document.core_properties.author = "STCR"

    # editorial_cover pattern, without the two cover blocks removed by request.
    document.add_paragraph()
    document.add_paragraph()
    add_title(document, "คู่มือการใช้งาน", size=28, after=2)
    add_title(document, "ระบบควบคุมอุณหภูมิเตารมควัน", size=24, after=8)
    add_title(document, "STCR — Smoking Temperature Control", size=13, color="F07818", after=18)
    add_body(
        document,
        "คู่มือสำหรับผู้ใช้งานบริษัท Grand Rubber (GR) และ TTN",
        size=11,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        bold=True,
    )
    add_body(
        document,
        "อธิบายหน้าหลักด้วยภาพจากระบบ พร้อมกรอบสีแดงและหมายเลขอ้างอิง",
        size=9.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=20,
    )

    toc_title = document.add_paragraph(style="Heading 1")
    set_run_font(toc_title.add_run("เนื้อหา"), size=16, bold=True, color=NAVY)
    toc_items = [
        "1. เข้าสู่ระบบ",
        "2. Dashboard และเมนูหลัก",
        "3. รายละเอียดเตาและข้อมูลเรียลไทม์",
        "4. การอ่านกราฟ",
        "5. ข้อมูลย้อนหลัง",
        "6. Alarm",
        "7. เลือกรายงานและดาวน์โหลดไฟล์",
        "8. กรอกข้อมูลฟอร์มรายงาน",
        "9. รายละเอียดรอบและข้อมูลเอกสาร",
        "10. Setting",
        "11. การแก้ปัญหาเบื้องต้น",
    ]
    table = document.add_table(rows=6, cols=2)
    set_table_geometry(table, [5103, 5102])
    for index, item in enumerate(toc_items):
        cell = table.cell(index % 6, index // 6)
        p = cell.paragraphs[0]
        set_paragraph_tokens(p, after=0, line=1.25)
        set_run_font(p.add_run(item), size=9.3, bold=True, color=INK)
    document.add_page_break()

    for section in manual.SECTIONS:
        add_section_page(document, section)
    add_troubleshooting(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    ARCHIVE.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    shutil.copy2(OUTPUT, ARCHIVE)
    print(OUTPUT)
    print(ARCHIVE)


if __name__ == "__main__":
    build_docx()
